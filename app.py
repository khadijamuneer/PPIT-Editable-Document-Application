import streamlit as st
from PIL import Image, ImageEnhance, ImageFilter, ImageOps
import easyocr
from docx import Document
from docx.shared import Pt
from io import BytesIO
import numpy as np
import cv2

# Configure page
st.set_page_config(
    page_title="Handwriting OCR Extractor",
    page_icon="✍️",
    layout="wide"
)

# Title
st.title("✍️ Handwriting OCR Text Extractor")
st.markdown("""
This app extracts text from **handwritten notes**. Due to the complexity of handwriting recognition, 
**manual review and correction is essential**. Use the split-screen editor to fix any OCR errors.
""")

# Sidebar
with st.sidebar:
    st.header("⚙️ OCR Settings")
    
    ocr_language = st.multiselect(
        "Language",
        ["en", "hi", "ar", "ur"],
        default=["en"],
        help="Select text languages"
    )
    
    st.header("🖼️ Image Enhancement")
    
    preprocessing_option = st.selectbox(
        "Preprocessing Mode",
        ["None", "Light Enhancement", "Strong Enhancement", "High Contrast", "Custom"],
        index=1
    )
    
    if preprocessing_option == "Custom":
        brightness_factor = st.slider("Brightness", 0.5, 2.0, 1.0, 0.1)
        contrast_factor = st.slider("Contrast", 0.5, 3.0, 1.5, 0.1)
        sharpness_factor = st.slider("Sharpness", 0.5, 3.0, 2.0, 0.1)
    
    st.header("📝 Output Settings")
    merge_lines = st.checkbox("Merge short lines", value=True, help="Combine lines that seem incomplete")
    remove_special_chars = st.checkbox("Clean special chars", value=False, help="Remove unusual characters")

# Load OCR
@st.cache_resource
def load_ocr(langs):
    return easyocr.Reader(langs, gpu=False)

def enhance_image_for_ocr(image, mode="Light Enhancement"):
    """Enhanced preprocessing specifically for handwritten text"""
    img = image.copy()
    
    # Convert to grayscale
    img = ImageOps.grayscale(img)
    
    if mode == "None":
        return img.convert('RGB')
    
    elif mode == "Light Enhancement":
        # Slight enhancement
        img = ImageEnhance.Contrast(img).enhance(1.3)
        img = ImageEnhance.Sharpness(img).enhance(1.5)
        img = ImageEnhance.Brightness(img).enhance(1.1)
    
    elif mode == "Strong Enhancement":
        # Stronger processing
        img = ImageEnhance.Contrast(img).enhance(2.0)
        img = ImageEnhance.Sharpness(img).enhance(2.5)
        # Apply bilateral filter to reduce noise but keep edges
        img_array = np.array(img)
        img_array = cv2.bilateralFilter(img_array, 9, 75, 75)
        img = Image.fromarray(img_array)
    
    elif mode == "High Contrast":
        # Maximum contrast for faded writing
        img = ImageEnhance.Contrast(img).enhance(2.5)
        img = ImageEnhance.Brightness(img).enhance(0.9)
        # Apply adaptive thresholding
        img_array = np.array(img)
        img_array = cv2.adaptiveThreshold(img_array, 255, cv2.ADAPTIVE_THRESH_GAUSSIAN_C, 
                                          cv2.THRESH_BINARY, 11, 2)
        img = Image.fromarray(img_array)
    
    elif mode == "Custom":
        img = ImageEnhance.Brightness(img).enhance(brightness_factor)
        img = ImageEnhance.Contrast(img).enhance(contrast_factor)
        img = ImageEnhance.Sharpness(img).enhance(sharpness_factor)
    
    return img.convert('RGB')

def extract_text_with_layout(image, reader):
    """Extract text while preserving layout as much as possible"""
    img_array = np.array(image)
    
    # Get detailed results with bounding boxes
    results = reader.readtext(img_array, detail=1, paragraph=False)
    
    if not results:
        return "⚠️ No text detected. Try adjusting image enhancement settings."
    
    # Sort by vertical position (top to bottom), then horizontal (left to right)
    sorted_results = sorted(results, key=lambda x: (x[0][0][1], x[0][0][0]))
    
    # Group into lines based on vertical proximity
    lines = []
    current_line = []
    current_y = None
    line_threshold = 30  # pixels
    
    for bbox, text, conf in sorted_results:
        y_pos = bbox[0][1]  # Top-left y coordinate
        
        if current_y is None or abs(y_pos - current_y) < line_threshold:
            current_line.append((bbox, text, conf))
            current_y = y_pos if current_y is None else current_y
        else:
            if current_line:
                lines.append(current_line)
            current_line = [(bbox, text, conf)]
            current_y = y_pos
    
    if current_line:
        lines.append(current_line)
    
    # Build text with proper spacing
    output_lines = []
    for line in lines:
        # Sort words in line by x position
        line_sorted = sorted(line, key=lambda x: x[0][0][0])
        line_text = ' '.join([text for _, text, _ in line_sorted])
        output_lines.append(line_text)
    
    return '\n'.join(output_lines)

def post_process_text(text, merge_short_lines=True, clean_chars=False):
    """Post-process extracted text"""
    if not text or text.startswith("⚠️"):
        return text
    
    lines = text.split('\n')
    
    # Merge short lines that might be incomplete
    if merge_short_lines:
        merged = []
        i = 0
        while i < len(lines):
            current = lines[i].strip()
            
            # If line is very short and next line exists, consider merging
            if len(current) < 40 and i + 1 < len(lines):
                next_line = lines[i + 1].strip()
                # Merge if current line doesn't end with punctuation
                if current and not current[-1] in '.!?:;':
                    merged.append(current + ' ' + next_line)
                    i += 2
                    continue
            
            merged.append(current)
            i += 1
        lines = merged
    
    # Clean special characters
    if clean_chars:
        lines = [line.replace('|', 'I').replace('~', '-') for line in lines]
    
    return '\n'.join(lines)

def create_document(texts, filenames):
    """Create formatted Word document"""
    doc = Document()
    
    # Title
    doc.add_heading('Extracted Handwritten Notes', 0)
    
    # Warning
    warning = doc.add_paragraph()
    warning.add_run('⚠️ IMPORTANT: ').bold = True
    warning.add_run('Handwriting OCR is not 100% accurate. Please review all text carefully and make necessary corrections.')
    
    # Content
    for i, (text, filename) in enumerate(zip(texts, filenames), 1):
        doc.add_heading(f'Page {i}: {filename}', 1)
        doc.add_paragraph(text)
        if i < len(texts):
            doc.add_page_break()
    
    # Save to bytes
    buffer = BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer

# Main app
uploaded_files = st.file_uploader(
    "📤 Upload Your Handwritten Notes",
    type=["png", "jpg", "jpeg", "bmp", "tiff"],
    accept_multiple_files=True,
    help="Upload clear photos or scans of your handwritten notes"
)

if uploaded_files:
    st.success(f"✅ {len(uploaded_files)} image(s) uploaded")
    
    # Load OCR
    with st.spinner("🔄 Loading OCR engine..."):
        try:
            reader = load_ocr(ocr_language)
        except Exception as e:
            st.error(f"❌ Error: {e}")
            st.stop()
    
    all_texts = []
    all_filenames = []
    
    # Process each image
    for idx, file in enumerate(uploaded_files):
        st.markdown(f"## 📄 Page {idx + 1}: {file.name}")
        
        col1, col2 = st.columns([1, 1])
        
        with col1:
            st.markdown("### 🖼️ Original Image")
            original = Image.open(file)
            st.image(original, use_container_width=True)
            
            # Show processed version
            with st.expander("🔍 View Processed Image"):
                processed = enhance_image_for_ocr(original, preprocessing_option)
                st.image(processed, use_container_width=True)
        
        with col2:
            st.markdown("### ✍️ Extracted Text (Editable)")
            
            with st.spinner("⏳ Extracting text..."):
                # Enhance and extract
                enhanced = enhance_image_for_ocr(original, preprocessing_option)
                raw_text = extract_text_with_layout(enhanced, reader)
                final_text = post_process_text(raw_text, merge_lines, remove_special_chars)
                
                # Editable text area
                edited = st.text_area(
                    "Review and correct the text below:",
                    final_text,
                    height=450,
                    key=f"text_{idx}",
                    help="⚠️ OCR may have errors - please review carefully!"
                )
                
                all_texts.append(edited)
                all_filenames.append(file.name)
                
                # Stats
                st.caption(f"📊 {len(edited.split())} words | {len(edited)} characters | {len(edited.splitlines())} lines")
        
        # Comparison tool
        with st.expander("🔄 Compare Original vs Extracted"):
            comp_col1, comp_col2 = st.columns(2)
            with comp_col1:
                st.image(original, caption="Original", use_container_width=True)
            with comp_col2:
                st.text_area("Extracted Text", edited, height=300, disabled=True, key=f"comp_{idx}")
        
        st.divider()
    
    # Download section
    st.markdown("## 📥 Download Your Text")
    
    download_col1, download_col2 = st.columns(2)
    
    with download_col1:
        # Word document
        doc_buffer = create_document(all_texts, all_filenames)
        st.download_button(
            "📄 Download as Word (.docx)",
            doc_buffer,
            "handwritten_notes.docx",
            "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            use_container_width=True,
            type="primary"
        )
    
    with download_col2:
        # Text file
        combined = "\n\n" + ("="*60 + "\n\n").join(
            [f"PAGE {i+1}: {name}\n{('='*60)}\n\n{text}" 
             for i, (name, text) in enumerate(zip(all_filenames, all_texts))]
        )
        st.download_button(
            "📝 Download as Text (.txt)",
            combined,
            "handwritten_notes.txt",
            "text/plain",
            use_container_width=True
        )

else:
    # Instructions
    st.info("👆 Upload images of your handwritten notes to get started")
    
    with st.expander("📖 How to Get Best Results"):
        st.markdown("""
        ### Tips for Better OCR Accuracy:
        
        #### 📸 Taking Photos:
        - ✅ Use good lighting (natural light is best)
        - ✅ Keep camera parallel to paper (avoid angles)
        - ✅ Ensure text is in focus
        - ✅ Use high resolution
        - ❌ Avoid shadows on the paper
        - ❌ Don't use flash if it creates glare
        
        #### ✍️ Writing Quality:
        - Clear, legible handwriting works best
        - Darker ink provides better contrast
        - Avoid very light or faded writing
        
        #### ⚙️ Using This App:
        1. Upload your image(s)
        2. Try different preprocessing modes if results aren't good
        3. **Always review and correct the extracted text**
        4. Download as Word or Text file
        
        ### ⚠️ Important Notes:
        - **Handwriting OCR is not perfect** - expect to make corrections
        - Mathematical symbols and special characters may not be recognized
        - Complex equations may need manual transcription
        - Chemical formulas often require correction
        """)

# Footer
st.markdown("---")
st.caption("💡 Tip: For best results with chemistry notes, you may need to manually correct chemical formulas and equations after extraction.")
